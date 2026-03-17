"""RevisedResumeSchema to DOCX output rendering.

Generates a clean, professional DOCX from structured resume data.
No server-side storage — generates on demand for download.
"""

import io
import logging

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import RevisedResumeSchema, ResumeSchema

logger = logging.getLogger(__name__)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def _add_contact_block(doc: Document, contact: dict) -> None:
    """Add contact information as a centered block at the top."""
    name = contact.get("name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)

    # Build contact line from available fields
    contact_parts = []
    for key in ["location", "email", "phone", "linkedin", "github"]:
        val = contact.get(key)
        if val:
            contact_parts.append(val)

    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.size = Pt(10)


def _add_section(doc: Document, title: str, sections: list, show_bullets: bool = True) -> None:
    """Add a resume section (experience, education, etc.)."""
    if not sections:
        return

    _add_heading(doc, title, level=2)

    for section in sections:
        # Section header (raw_text typically contains role/company/dates)
        p = doc.add_paragraph()
        run = p.add_run(section.raw_text)
        run.font.size = Pt(11)

        # Bullets
        if show_bullets and section.bullets:
            for bullet in section.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)


def render_docx(resume: RevisedResumeSchema | ResumeSchema) -> bytes:
    """Render a resume schema to DOCX bytes.

    Args:
        resume: Structured resume data (original or revised).

    Returns:
        DOCX file content as bytes, ready for download.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Contact block
    _add_contact_block(doc, resume.contact)

    # Summary
    if resume.summary:
        _add_heading(doc, "Professional Summary", level=2)
        p = doc.add_paragraph(resume.summary)
        p.paragraph_format.space_after = Pt(6)

    # Skills
    if resume.skills:
        _add_heading(doc, "Skills", level=2)
        p = doc.add_paragraph(", ".join(resume.skills))
        p.paragraph_format.space_after = Pt(6)

    # Experience
    _add_section(doc, "Experience", resume.experience)

    # Education
    _add_section(doc, "Education", resume.education, show_bullets=False)

    # Certifications
    if resume.certifications:
        _add_heading(doc, "Certifications", level=2)
        for cert in resume.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    logger.info("DOCX rendered: %d bytes", buf.getbuffer().nbytes)
    return buf.getvalue()
