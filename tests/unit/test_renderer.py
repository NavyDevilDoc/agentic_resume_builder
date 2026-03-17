"""Tests for DOCX renderer — smoke test and content verification."""

import io
from docx import Document

from models.schemas import ResumeSection, RevisedResumeSchema, ResumeSchema
from modules.output.renderer import render_docx


class TestRenderDocx:
    def _sample_revised(self) -> RevisedResumeSchema:
        return RevisedResumeSchema(
            contact={"name": "Jeremy Springston", "email": "jeremy@example.com", "location": "Fredericksburg, VA"},
            summary="AI/ML Engineer with 20+ years of experience.",
            experience=[
                ResumeSection(
                    section_type="experience",
                    raw_text="Booz Allen Hamilton, AI/ML Engineer 3, 05/2025 – Present",
                    bullets=[
                        "Deployed semantic search on AWS infrastructure",
                        "Built RAG pipeline with Python and LangChain",
                    ],
                )
            ],
            skills=["Python", "PyTorch", "AWS", "Docker"],
            education=[
                ResumeSection(
                    section_type="education",
                    raw_text="MS Johns Hopkins University, Applied Mathematics 2021",
                )
            ],
            certifications=["AWS Solutions Architect"],
            rewrite_level_used="edit",
            change_log=["Added AWS to skills"],
        )

    def test_returns_bytes(self):
        result = render_docx(self._sample_revised())
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx(self):
        """Output bytes are a valid DOCX that python-docx can read."""
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_contains_name(self):
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jeremy Springston" in all_text

    def test_contains_experience(self):
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Booz Allen Hamilton" in all_text
        assert "semantic search" in all_text

    def test_contains_skills(self):
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Python" in all_text
        assert "AWS" in all_text

    def test_contains_education(self):
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Johns Hopkins" in all_text

    def test_contains_certifications(self):
        result = render_docx(self._sample_revised())
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AWS Solutions Architect" in all_text

    def test_empty_resume(self):
        """Minimal resume should still render without error."""
        minimal = RevisedResumeSchema(
            contact={},
            rewrite_level_used="edit",
            change_log=[],
        )
        result = render_docx(minimal)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_original_resume_schema(self):
        """Renderer should also work with non-revised ResumeSchema."""
        original = ResumeSchema(
            contact={"name": "Test User"},
            summary="A summary.",
            skills=["Python"],
        )
        result = render_docx(original)
        doc = Document(io.BytesIO(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Test User" in all_text
